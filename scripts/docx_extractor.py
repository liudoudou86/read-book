#!/usr/bin/env python3
"""
DOCX 书籍解析工具
回退链：python-docx → stdlib zipfile+xml
"""

import argparse
import xml.etree.ElementTree as ET
from pathlib import Path
from zipfile import ZipFile


def extract_pythondocx(docx_path):
    """方式1: python-docx"""
    try:
        import docx
        doc = docx.Document(docx_path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        # 也提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception:
        pass
    return None


def extract_zipfallback(docx_path):
    """方式2: stdlib zipfile + XML 回退"""
    try:
        text_parts = []
        with ZipFile(docx_path) as z:
            # 读取 document.xml
            if "word/document.xml" not in z.namelist():
                return None
            xml_content = z.read("word/document.xml")
            root = ET.fromstring(xml_content)
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            for t_elem in root.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                if t_elem.text:
                    text_parts.append(t_elem.text)
                # tail text
                if t_elem.tail:
                    text_parts.append(t_elem.tail)
        text = "".join(text_parts)
        import re
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        return text if text else None
    except Exception:
        pass
    return None


def extract_text_from_docx(docx_path, output_path=None):
    docx_path = Path(docx_path)
    if not docx_path.exists():
        print(f"文件不存在：{docx_path}")
        return None

    print(f"正在读取 DOCX：{docx_path}")

    text = None
    for name, extractor in [("python-docx", extract_pythondocx), ("zipfile 回退", extract_zipfallback)]:
        print(f"  尝试：{name}...")
        result = extractor(docx_path)
        if result and result.strip():
            text = result
            print(f"  成功：{name}")
            break

    if not text:
        print("所有 DOCX 引擎均提取失败")
        print("提示：安装 python-docx 可提升提取质量")
        return None

    if output_path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(text, encoding="utf-8")
        print(f"已保存到：{output_path}")
    else:
        output_path = docx_path.with_suffix(".txt")
        output_path.write_text(text, encoding="utf-8")
        print(f"已保存到：{output_path}")

    print(f"  字符数：{len(text)}")
    return output_path


def main():
    parser = argparse.ArgumentParser(description="DOCX 书籍解析工具")
    parser.add_argument("--file", "-f", type=str, required=True, help="DOCX 文件路径")
    parser.add_argument("--output", "-o", type=str, help="输出文本文件路径")
    args = parser.parse_args()
    result = extract_text_from_docx(args.file, args.output)
    if result:
        print("\n下一步：使用 Read 工具读取生成的文本文件")


if __name__ == "__main__":
    main()
